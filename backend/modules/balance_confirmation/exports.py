"""Word/PDF generators for the Balance Confirmation utility.

For Phase 1+2 we ship:
  • build_authorization_template_docx — editable Word letter the client
    signs and re-uploads as the standing authorisation PDF.

PDF letter generation per-confirmation will land in Phase 3 alongside the
sending engine.
"""
from __future__ import annotations
import io
from datetime import datetime, timezone
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_authorization_template_docx(client: Dict[str, Any]) -> bytes:
    """Generate an editable Word template the client can fill on their letterhead.

    The CA passes this to the client; the client signs, scans as PDF, and
    re-uploads via the Authorization upload endpoint. That signed PDF is then
    auto-attached to every confirmation email under Phase 3.
    """
    doc = Document()

    # Page margins kept default (1 inch).
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AUTHORISATION LETTER")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x04, 0x78, 0x57)  # AssureAI green

    doc.add_paragraph()  # spacing

    # Top metadata block
    meta = doc.add_paragraph()
    meta.add_run(f"From: {(client or {}).get('name', '[Client Name]')}").bold = True
    meta.add_run(
        f"\nGSTIN: {(client or {}).get('gstin') or '[GSTIN]'}"
        f"\nFile No: {(client or {}).get('file_number') or '[File No]'}"
        f"\n\nDate: ____________________"
        f"\nTo,\nThe Branch Manager / Authorised Person\n[Bank / Vendor / Customer Name]\n[Address]"
    )

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Sub: Authorisation to share information directly with our auditors")
    sub_run.bold = True

    doc.add_paragraph()

    body = doc.add_paragraph(
        "Dear Sir / Madam,"
    )
    body.add_run(
        "\n\nIn connection with the statutory audit of our books of account "
        "for the financial year ended __________, we hereby authorise you to "
        "provide our Statutory Auditors, M/s [Auditor Firm Name], directly "
        "with the following information as on the said date:"
    )

    bullets = [
        "Balance outstanding in our account(s) with you, including any "
        "current, term, fixed deposit, overdraft, cash credit or loan accounts.",
        "Details of any guarantees, letters of credit, hypothecations, "
        "mortgages, liens or other charges in our favour or against us.",
        "Statement of account / ledger extract for the financial year, "
        "showing all debits and credits, contra entries and balances.",
        "Particulars of any unbilled invoices, advances or amounts in transit.",
        "Any other information our auditors may reasonably require for the "
        "purpose of audit verification.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph(
        "\nThe response may be sent directly to our auditors at the email "
        "address from which their request originates, or by clicking the "
        "secure confirmation link contained in their email."
    )
    doc.add_paragraph(
        "We confirm that this authorisation supersedes any earlier instruction "
        "and shall remain valid until the completion of the audit."
    )

    doc.add_paragraph()
    doc.add_paragraph("Thanking you,")
    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph("For " + ((client or {}).get("name") or "[Client Name]"))
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("(Authorised Signatory)")
    doc.add_paragraph("Name: ____________________")
    doc.add_paragraph("Designation: ______________")

    # Footer note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = p.add_run(
        f"\n\nGenerated {datetime.now(timezone.utc).strftime('%d %b %Y')} · "
        f"AssureAI Audit Utilities · Sign on company letterhead, "
        f"then re-upload the signed PDF."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
